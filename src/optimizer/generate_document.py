from docx import Document
from datetime import datetime

def generate_doc(route):
    document = Document()
    document.add_heading("Move Itinerary", 0)
    document.add_paragraph("Generated on {}".format(datetime.now()))

    document.add_heading("Summary", 2)
    document.add_paragraph("Total distance: {:.2f} km".format(route.get_total_distance()/1000))
    duration = route.get_total_duration()
    durationHours = duration//(60*60)
    durationMin = (duration - durationHours*60*60)/60
    document.add_paragraph("Total duration: {} h {:.2f} min".format(durationHours, durationMin))

    document.add_heading("Detailed Itinerary", 2)
    detailedRoute = route.get_route()
    table = document.add_table(rows = 1, cols=6)
    headers = ['Route Leg #', "Expected Time (m)", "Total time (m)",  "Starting Address", "Ending Address", "Distance (m)"]
    row_populator(table.rows[0].cells, headers)
    time = 0
    for index, leg in enumerate(detailedRoute):
        rowCells = table.add_row().cells
        rowCells[0].text = str(index+1)
        rowCells[1].text = "{:.2f}".format(leg['duration']/60)
        time += leg['duration']/60
        rowCells[2].text = "{:.2f}".format(time)
        rowCells[3].text = "{}\n ({})".format(leg['start'].get_address(), leg['start'].get_role())
        rowCells[4].text = "{}\n ({})".format(leg['end'].get_address(), leg['end'].get_role())
        rowCells[5].text = "{} m".format(leg['distance'])
    
    docName = 'Move_itinerary_{}.docx'.format(datetime.now().date())
    document.save(docName)
    return docName

def row_populator(rowCells, content):
    for index, val in enumerate(content):
        rowCells[index].text = val


