import unittest
from unittest.mock import patch

import csv
import os 
import datetime
import docx

from optimizer.route_creator import RouteCreator
from optimizer.generate_document import generate_doc

class TestWholeSystem(unittest.TestCase):

    def setUp(self):
        self.create_file()
    
    def tearDown(self):
        self.delete_csv_file()
        self.delete_word_doc()
    
    def test_doc_created_with_correct_name(self):
        r = RouteCreator('locations.csv')
        docName = generate_doc(r)
        self.assertEqual(docName, 'Move_itinerary_{}.docx'.format(datetime.datetime.now().date()))
    
    def test_doc_created_with_correct_table_and_summary(self):
        r = RouteCreator('locations.csv')
        doc = docx.Document(generate_doc(r))
        content = [p.text for p in doc.paragraphs]
        headers = [cell.text for cell in doc.tables[0].rows[0].cells]
        firstTableRow = [cell.text for cell in doc.tables[0].rows[1].cells]
        secondTableRow = [cell.text for cell in doc.tables[0].rows[2].cells]
        self.assertIn('Total distance: 2.91 km', content)
        self.assertIn('Total duration: 0.0 h 7.21 min', content)
        self.assertIn('Move Itinerary', content)
        self.assertEqual(headers, ['Route Leg #', 'Expected Time (m)', 'Total time (m)', 'Starting Address', 'Ending Address', 'Distance (m)'])
        self.assertEqual(firstTableRow, ['1', '3.63', '3.63', '431 College St, Toronto\n (car)', '10 Kings College Rd, Toronto\n (volunteer)', '1452.2 m'])
        self.assertEqual(secondTableRow, ['2', '3.58', '7.21', '10 Kings College Rd, Toronto\n (volunteer)', '431 College St, Toronto\n (car)', '1452.9 m'])

    def test_lookup_error_thrown_if_address_cannot_be_converted_to_coord(self):
        self.create_file_bad_address()
        with self.assertRaises(LookupError) as context:
            RouteCreator('locationsbad.csv')
        self.delete_bad_address_file()
        print(context.exception)
        self.assertTrue('Coordinate lookup failed for 567 Hawaii Fake, Toronto. Check if valid address.' in str(context.exception))


    def create_file(self):
        with open('locations.csv', 'w') as csvfile:
            filewriter = csv.writer(csvfile, delimiter=',', quotechar='|', quoting=csv.QUOTE_MINIMAL)
            filewriter.writerow(['street','city','country','postalcode','mandatory','role','routeRestriction','nearestIntersection'])
            filewriter.writerow(['10 Kings College Rd', 'Toronto', 'Canada', 'M5S3G8', 'TRUE', 'volunteer', 'end'])
            filewriter.writerow(['431 College St', 'Toronto', 'Canada', 'M5T1T1', 'TRUE', 'car', 'start'])
    
    def create_file_bad_address(self):
        with open('locationsbad.csv', 'w') as csvfile:
            filewriter = csv.writer(csvfile, delimiter=',', quotechar='|', quoting=csv.QUOTE_MINIMAL)
            filewriter.writerow(['street','city','country','postalcode','mandatory','role','routeRestriction','nearestIntersection'])
            filewriter.writerow(['567 Hawaii Fake', 'Toronto', 'Canada', 'M5S3G8', 'TRUE', 'volunteer', 'end'])
            filewriter.writerow(['431 College St', 'Toronto', 'Canada', 'M5T1T1', 'TRUE', 'car', 'start'])
    
    def delete_bad_address_file(self):
        if os.path.isfile('locationsbad.csv'):
            os.remove('locationsbad.csv')

    def delete_csv_file(self):
        if os.path.isfile('locations.csv'):
            os.remove('locations.csv')

    def delete_word_doc(self):
        filename = 'Move_itinerary_{}.docx'.format(datetime.datetime.now().date())
        if os.path.isfile(filename):
            os.remove(filename)

    